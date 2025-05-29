from typing import Optional, List, Dict, Any
import os
from pathlib import Path
import PyPDF2
from docx import Document
import re

class DocumentProcessor:
    def __init__(self):
        """Inicializa o processador de documentos."""
        self.supported_extensions = {'.pdf', '.docx', '.txt'}
        
    def read_document(self, file_path: str) -> str:
        """
        Lê o conteúdo de um documento.
        
        Args:
            file_path (str): Caminho do arquivo
            
        Returns:
            str: Conteúdo do documento
            
        Raises:
            ValueError: Se o formato do arquivo não for suportado
            FileNotFoundError: Se o arquivo não existir
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Arquivo não encontrado: {file_path}")
            
        if file_path.suffix not in self.supported_extensions:
            raise ValueError(f"Formato não suportado: {file_path.suffix}")
            
        if file_path.suffix == '.pdf':
            return self._read_pdf(file_path)
        elif file_path.suffix == '.docx':
            return self._read_docx(file_path)
        else:  # .txt
            return self._read_txt(file_path)
            
    def _read_pdf(self, file_path: Path) -> str:
        """Lê o conteúdo de um arquivo PDF."""
        text = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text.append(page.extract_text())
        return '\n'.join(text)
    
    def _read_docx(self, file_path: Path) -> str:
        """Lê o conteúdo de um arquivo DOCX."""
        doc = Document(file_path)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    
    def _read_txt(self, file_path: Path) -> str:
        """Lê o conteúdo de um arquivo de texto."""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
            
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extrai metadados do documento.
        
        Args:
            file_path (str): Caminho do arquivo
            
        Returns:
            Dict[str, Any]: Metadados do documento
        """
        file_path = Path(file_path)
        metadata = {
            'filename': file_path.name,
            'extension': file_path.suffix,
            'size': os.path.getsize(file_path),
            'created': os.path.getctime(file_path),
            'modified': os.path.getmtime(file_path)
        }
        
        if file_path.suffix == '.pdf':
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata.update({
                    'pages': len(pdf_reader.pages),
                    'author': pdf_reader.metadata.get('/Author', ''),
                    'title': pdf_reader.metadata.get('/Title', ''),
                    'subject': pdf_reader.metadata.get('/Subject', '')
                })
        elif file_path.suffix == '.docx':
            doc = Document(file_path)
            metadata.update({
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'sections': len(doc.sections)
            })
            
        return metadata
        
    def extract_tables(self, file_path: str) -> List[List[List[str]]]:
        """
        Extrai tabelas do documento.
        
        Args:
            file_path (str): Caminho do arquivo
            
        Returns:
            List[List[List[str]]]: Lista de tabelas extraídas
        """
        file_path = Path(file_path)
        tables = []
        
        if file_path.suffix == '.docx':
            doc = Document(file_path)
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                tables.append(table_data)
                
        return tables
        
    def extract_images(self, file_path: str, output_dir: Optional[str] = None) -> List[str]:
        """
        Extrai imagens do documento.
        
        Args:
            file_path (str): Caminho do arquivo
            output_dir (Optional[str]): Diretório para salvar as imagens
            
        Returns:
            List[str]: Lista de caminhos das imagens extraídas
        """
        # TODO: Implementar extração de imagens
        return [] 